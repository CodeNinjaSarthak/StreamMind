"""RAG document upload, chunking, and retrieval service."""

import io
import logging
import re
from typing import List
from uuid import UUID

import docx
import pypdf
from app.db.models.rag import RAGDocument
from app.services.gemini.client import (
    GeminiClient,
    vector_to_literal,
)
from fastapi import (
    HTTPException,
    UploadFile,
)
from sqlalchemy import text
from sqlalchemy.orm import Session

logger = logging.getLogger(__name__)

MAX_FILE_SIZE = 10 * 1024 * 1024  # 10 MB
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}
CHUNK_SIZE_WORDS = 500
CHUNK_OVERLAP_WORDS = 100


def _extract_text_pdf(content: bytes) -> str:
    reader = pypdf.PdfReader(io.BytesIO(content))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _extract_text_docx(content: bytes) -> str:
    doc = docx.Document(io.BytesIO(content))
    return "\n".join(para.text for para in doc.paragraphs)


def _chunk_text(full_text: str) -> List[str]:
    """Split text into overlapping word-based chunks using sentence boundaries."""
    sentences = re.split(r"(?<=[.!?])\s+", full_text)
    chunks = []
    current_words: List[str] = []

    for sentence in sentences:
        sentence_words = sentence.split()
        current_words.extend(sentence_words)
        if len(current_words) >= CHUNK_SIZE_WORDS:
            chunks.append(" ".join(current_words))
            # Keep tail for overlap
            current_words = current_words[-CHUNK_OVERLAP_WORDS:]

    if current_words:
        chunks.append(" ".join(current_words))

    return [c for c in chunks if c.strip()]


async def upload_document(
    file: UploadFile,
    teacher_id: UUID,
    db: Session,
) -> List[RAGDocument]:
    """Upload and chunk a document, store embeddings in DB.

    Args:
        file: Uploaded file (PDF or DOCX).
        teacher_id: Owning teacher's UUID.
        db: SQLAlchemy session.

    Returns:
        List of created RAGDocument records.
    """
    if file.content_type not in ALLOWED_MIME_TYPES:
        raise HTTPException(status_code=415, detail=f"Unsupported file type: {file.content_type}")

    content = await file.read()
    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(status_code=413, detail="File exceeds 10 MB limit")

    # Magic byte validation
    if file.content_type == "application/pdf" and not content.startswith(b"%PDF"):
        raise HTTPException(status_code=400, detail="Invalid PDF file")

    if file.content_type == "application/pdf":
        full_text = _extract_text_pdf(content)
        source_type = "pdf"
    else:
        full_text = _extract_text_docx(content)
        source_type = "docx"

    chunks = _chunk_text(full_text)
    if not chunks:
        raise HTTPException(status_code=422, detail="Could not extract any text from document")

    gemini = GeminiClient()
    records = []
    base_title = file.filename or "document"

    for i, chunk in enumerate(chunks):
        try:
            embedding = gemini.generate_embedding(chunk)
        except Exception as e:
            logger.error(f"Embedding generation failed at chunk {i + 1}/{len(chunks)}: {e}")
            db.rollback()
            raise HTTPException(status_code=502, detail="Embedding generation failed. Please try again.")
        doc = RAGDocument(
            teacher_id=teacher_id,
            title=f"{base_title} (chunk {i + 1}/{len(chunks)})",
            content=chunk,
            source_type=source_type,
        )
        doc.embedding = embedding
        db.add(doc)
        records.append(doc)

    db.commit()
    logger.info(f"Uploaded {len(records)} chunks for teacher {teacher_id} from '{base_title}'")
    return records


def search_similar(
    query_embedding: list,
    teacher_id: UUID,
    db: Session,
    limit: int = 5,
) -> list:
    """Find RAG documents similar to the given embedding."""
    rows = db.execute(
        text(
            "SELECT id, title, content FROM rag_documents "
            "WHERE teacher_id = CAST(:tid AS uuid) "
            "ORDER BY embedding <-> CAST(:emb AS vector) LIMIT :lim"
        ),
        {"tid": str(teacher_id), "emb": vector_to_literal(query_embedding), "lim": limit},
    ).fetchall()
    return rows
